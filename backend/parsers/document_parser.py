"""
Document parser for PDF and DOCX files with chunking.
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

from docx import Document
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger("kms.parser")


class DocumentParser:
    """Parses PDF and DOCX documents into chunks."""

    CHUNK_SIZE = 512
    CHUNK_OVERLAP = 80

    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ):
        self.chunk_size = chunk_size if chunk_size is not None else self.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap if chunk_overlap is not None else self.CHUNK_OVERLAP
        self._text_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="cl100k_base",
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=[
                "\n\n", "\n", # Note that "\n\n" may NOT represent a section separator in the text parsed from pdf/docx 
                "。", "！", "？",
                "；", "，", "、",
                " ",
                ""],
        )

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a document and return chunks.
        Supports PDF and DOCX files.
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == ".pdf":
            return self._parse_pdf(file_path)
        elif extension == ".docx":
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse PDF file and extract text with page info."""
        chunks = []
        reader = PdfReader(file_path)

        full_text = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                full_text.append(f"[Page {page_num + 1}]\n{text}")

        combined_text = "\n\n".join(full_text)
        text_chunks = self._chunk_text(combined_text)

        for i, chunk_text in enumerate(text_chunks, 1):
            # Determine which page this chunk is from
            page_match = re.search(r'\[Page (\d+)\]', chunk_text)
            page = int(page_match.group(1)) if page_match else 1

            # Clean up page markers from content
            clean_content = re.sub(r'\[Page \d+\]\n?', '', chunk_text).strip()

            chunks.append({
                "content": clean_content,
                "metadata": {
                    "page": page,
                    "chunk_index": i,
                    "source": "pdf"
                }
            })

        return chunks

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Parse DOCX file and extract text with paragraph info."""
        doc = Document(file_path)
        chunks = []

        # Extract paragraphs with their indices
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append((i, text))

        # Build full text and chunk
        full_text = "\n\n".join([f"[Para {idx}] {text}" for idx, text in paragraphs])
        text_chunks = self._chunk_text(full_text)

        for i, chunk_text in enumerate(text_chunks):
            # Determine paragraph range
            para_matches = re.findall(r'\[Para (\d+)\]', chunk_text)
            if para_matches:
                start_para = min(int(p) for p in para_matches)
                end_para = max(int(p) for p in para_matches)
                para_range = f"{start_para}-{end_para}"
            else:
                para_range = "unknown"

            # Clean up markers
            clean_content = re.sub(r'\[Para \d+\]', '', chunk_text).strip()

            chunks.append({
                "content": clean_content,
                "metadata": {
                    "paragraph_range": para_range,
                    "chunk_index": i,
                    "source": "docx"
                }
            })

        return chunks

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into fixed-size token chunks with overlap."""
        text_chunks = self._text_splitter.split_text(text)
        logger.info(f"Parsed into {len(text_chunks)} chunks")
        return text_chunks

    def _markdown_to_text(self, markdown: str) -> str:
        """Strip markdown syntax to plain text."""
        text = re.sub(r"!\[.*?\]\(.*?\)", "", markdown)
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"__(.+?)__", r"\1", text)
        text = re.sub(r"_(.+?)_", r"\1", text)
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"^[-*_]{3,}$", "", text, flags=re.MULTILINE)
        return text.strip()

    def _add_page_markers(self, text: str) -> str:
        """Detect page boundaries, mark with [Page N]."""
        if re.search(r"---.*?---", text, re.DOTALL):
            pages = re.split(r"(?=\n---\n)", text)
            if len(pages) > 1:
                return "\n\n".join(
                    f"[Page {i}]\n{page.strip()}"
                    for i, page in enumerate(pages, 1)
                )
        return f"[Page 1]\n{text}"

    def _add_paragraph_markers(self, text: str) -> str:
        """Split text into paragraphs, mark with [Para N]."""
        paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]
        if not paragraphs:
            return text
        return "\n\n".join(
            f"[Para {i}]\n{p}" for i, p in enumerate(paragraphs, 1)
        )

    def _build_chunks(
        self, markdown_content: str, file_extension: str
    ) -> List[Dict[str, Any]]:
        """Parse markdown, add markers, chunk, and return in DocumentParser format."""
        plain_text = self._markdown_to_text(markdown_content)

        if file_extension == ".pdf":
            marked_text = self._add_page_markers(plain_text)
        else:
            marked_text = self._add_paragraph_markers(plain_text)

        text_chunks = self._chunk_text(marked_text)

        chunks = []
        for i, chunk_text in enumerate(text_chunks, 1):
            if file_extension == ".pdf":
                page_match = re.search(r"\[Page (\d+)\]", chunk_text)
                page = int(page_match.group(1)) if page_match else 1
                clean_content = re.sub(r"\[Page \d+\]\n?", "", chunk_text).strip()
                chunks.append({
                    "content": clean_content,
                    "metadata": {"page": page, "chunk_index": i, "source": "pdf"}
                })
            else:
                para_matches = re.findall(r"\[Para (\d+)\]", chunk_text)
                if para_matches:
                    start_para = min(int(p) for p in para_matches)
                    end_para = max(int(p) for p in para_matches)
                    para_range = f"{start_para}-{end_para}"
                else:
                    para_range = "unknown"
                clean_content = re.sub(r"\[Para \d+\]", "", chunk_text).strip()
                chunks.append({
                    "content": clean_content,
                    "metadata": {
                        "paragraph_range": para_range,
                        "chunk_index": i,
                        "source": "docx"
                    }
                })
        return chunks
